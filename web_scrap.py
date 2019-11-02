import docx
import requests
import urllib.request
from bs4 import BeautifulSoup

# scraping

search = str(input("Enter the Search item:").capitalize())
query = search.split()
list=[]
for q in query:
    query1 = q.capitalize()
    list.append(query1)
# print(list)

con = "_".join(list)
img = "-".join(list)

r = requests.get("https://en.wikipedia.org/wiki/"+con)
page = r.content

soup = BeautifulSoup(page, "html.parser")
heading = soup.find("h1", {"class": "firstHeading"})

head = heading.text
# getting image from internet
i = requests.get("https://www.gettyimages.in/photos/"+img)
page1 = i.content
soup1 = BeautifulSoup(page1, "html.parser")
image = soup1.find_all("img", {"class": "gallery-asset__thumb gallery-mosaic-asset__thumb"})
url = image[1].get('src')
file_name = search
full_file_path = file_name+".jpg"


def dl_jpg(url, file_path, file_name):
    full_path = file_path+file_name+'.jpg'
    urllib.request.urlretrieve(url, full_path)


dl_jpg(url, 'images/', file_name)

# getting content from internet
body = soup.find_all("p", {})
content_list = []
for i in range(len(body)):
    content_list.append(body[i].text)
full_content = str("\n".join(content_list))

# writing into a word file
doc = docx.Document()
doc.add_heading(head, 0)
doc.add_picture('images/'+full_file_path, width=docx.shared.Inches(6), height=docx.shared.Inches(4))
doc.add_paragraph(full_content[:6000])
doc.save('documents/{}.docx'.format(con))

print("Completed the Assignment !!!!")
print("File Generated at: D:\Python3\Assignment Maker\documents")